import streamlit as st
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import os
import json
from openai import OpenAI
import re # Dodano import re
#import logging # Dodano import logging
import base64
from datetime import datetime

import markdown
from docx import Document
from docx.shared import Pt, Inches
import tempfile
from weasyprint import HTML, CSS
from bs4 import BeautifulSoup

# --- Konfiguracja logowania ---
#logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Konfiguracja ---

# WAŻNE: Wstaw tutaj swój klucz API OpenRouter.
# Pamiętaj, aby nie umieszczać prawdziwego klucza w publicznych repozytoriach!
# Możesz też wczytywać go ze zmiennej środowiskowej dla większego bezpieczeństwa.
OPENROUTER_API_KEY = st.secrets["Openrouter_key"]  # <--- ZASTĄP SWOIM KLUCZEM

# Wczytaj komunikaty systemowe z pliku JSON
try:
    with open("system_messages.json", "r", encoding="utf-8") as f:
        SYSTEM_MESSAGES = json.load(f)
except FileNotFoundError:
    st.error("Nie znaleziono pliku system_messages.json!")
    SYSTEM_MESSAGES = {"default": "Error: system_messages.json not found."}
except json.JSONDecodeError:
    st.error("Błąd podczas parsowania pliku system_messages.json!")
    SYSTEM_MESSAGES = {"default": "Error: Could not parse system_messages.json."}

# Opcjonalna wiadomość systemowa dla modelu LLM. Można ją dostosować.
# SYSTEM_MESSAGE = """... (usunięto długi ciąg znaków) ...""" # Usunięto statyczną definicję

# Dostępne języki (Wyświetlana nazwa: kod Tesseract / kod dla LLM)
LANGUAGES = {
    "Angielski": ("eng", "English"),
    "Niemiecki": ("deu", "German"),
    "Polski": ("pol", "Polish"),
    "Gruziński": ("kat", "Georgian"),
    "Ukraiński": ("ukr", "Ukrainian"),
    "Chiński (Uproszczony)": ("chi-sim", "Simplified Chinese"),
    "Chiński (Tradycyjny)": ("chi-tra", "Traditional Chinese"),
    "Francuski": ("fra", "French"),
    "Hiszpański": ("spa", "Spanish"),
    "Hinduski": ("hin", "Hindi"),
    "Turecki": ("tur", "Turkish"),
}

# --- Funkcje pomocnicze ---

def parse_page_numbers(selection_str, max_pages):
    """Parsuje ciąg znaków z numerami stron (np. '1, 3, 5-7') do listy."""
    pages = set()
    if not selection_str or not selection_str.strip():
        # Jeśli puste, zwróć wszystkie strony
        return list(range(1, max_pages + 1))

    try:
        # Usuń białe znaki i podziel po przecinkach
        parts = selection_str.replace(" ", "").split(',')
        for part in parts:
            if not part: continue # Ignoruj puste części (np. po podwójnym przecinku)
            if '-' in part:
                # Obsługa zakresu
                start_str, end_str = part.split('-')
                start = int(start_str)
                end = int(end_str)
                if not (1 <= start <= end <= max_pages):
                    raise ValueError(f"Nieprawidłowy zakres stron: {part}. Dostępne: 1-{max_pages}")
                pages.update(range(start, end + 1))
            else:
                # Obsługa pojedynczej strony
                page_num = int(part)
                if not (1 <= page_num <= max_pages):
                    raise ValueError(f"Numer strony poza zakresem: {page_num}. Dostępne: 1-{max_pages}")
                pages.add(page_num)

        if not pages:
             raise ValueError("Nie wybrano żadnych stron.")

        return sorted(list(pages))
    except ValueError as e:
        st.error(f"Błąd w formacie wyboru stron ('{selection_str}'): {e}")
        return None # Zwróć None w przypadku błędu

def extract_images_from_pdf(pdf_bytes, selected_pages: list[int] | None = None):
    """Ekstrahuje obrazy wybranych stron z pliku PDF."""
    images = []
    doc = None # Inicjalizacja doc
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages_in_doc = len(doc)

        # Jeśli nie podano wybranych stron, użyj wszystkich
        if selected_pages is None:
            pages_to_process = list(range(1, total_pages_in_doc + 1))
        else:
            # Walidacja, czy wybrane strony istnieją w dokumencie
            invalid_pages = [p for p in selected_pages if not (1 <= p <= total_pages_in_doc)]
            if invalid_pages:
                st.error(f"Numery stron poza zakresem: {invalid_pages}. Dokument ma {total_pages_in_doc} stron.")
                return None, total_pages_in_doc # Zwracamy też liczbę stron dla informacji
            pages_to_process = selected_pages

        if not pages_to_process:
             st.warning("Nie wybrano żadnych stron do przetworzenia.")
             return [], total_pages_in_doc

        # Iteruj tylko po wybranych (i zwalidowanych) numerach stron (1-indeksowane)
        for page_num_one_indexed in pages_to_process:
            page_index_zero_based = page_num_one_indexed - 1 # Konwersja na 0-indeksowanie
            page = doc.load_page(page_index_zero_based)
            # Renderuj stronę jako obraz PNG w wysokiej rozdzielczości (300 DPI)
            pix = page.get_pixmap(dpi=300)
            img_bytes = pix.tobytes("png")
            images.append(Image.open(io.BytesIO(img_bytes)))

        return images, total_pages_in_doc
    except Exception as e:
        st.error(f"Błąd podczas przetwarzania PDF: {e}")
        return None, 0 # Zwróć 0 stron w razie błędu otwierania
    finally:
        if doc:
            doc.close() # Upewnij się, że dokument jest zamknięty

def perform_ocr(images, lang_code):
    """Wykonuje OCR na liście obrazów używając Tesseract, z automatycznym wykrywaniem orientacji."""
    full_text = ""
    if not images: # Dodano sprawdzenie czy lista obrazów nie jest pusta
        return ""

    try:
        # Sprawdzenie, czy Tesseract jest zainstalowany i dostępny
        try:
            pytesseract.get_tesseract_version()
        except pytesseract.TesseractNotFoundError:
            st.error("Tesseract nie jest zainstalowany lub nie ma go w ścieżce systemowej (PATH).")
            st.error("Instrukcje instalacji: https://tesseract-ocr.github.io/tessdoc/Installation.html")
            return None

        #st.write("Rozpoczynanie OCR z wykrywaniem orientacji...") # Dodano komunikat
        for i, img in enumerate(images):
            st.write(f"Przetwarzanie strony {i+1}/{len(images)}")
            rotated_image = img # Domyślnie użyj oryginalnego obrazu

            try:
                # 1. Wykryj orientację i skrypt
                # Używamy --psm 0 dla OSD
                osd_data = pytesseract.image_to_osd(img, output_type=pytesseract.Output.DICT, config='--psm 0')
                rotation = osd_data.get('rotate', 0) # Pobierz kąt rotacji (domyślnie 0)
                script = osd_data.get('script', 'N/A') # Pobierz wykryty skrypt

                #st.write(f"  > Strona {i+1}: Wykryto skrypt: {script}, Rotacja: {rotation}°")

                # 2. Obróć obraz, jeśli to konieczne
                if rotation != 0:
                    #st.write(f"  > Strona {i+1}: Korygowanie rotacji o {-rotation}°...")
                    # Obracamy w przeciwnym kierunku niż wykryta rotacja
                    # expand=True zapobiega przycinaniu obrazu
                    rotated_image = img.rotate(-rotation, resample=Image.BICUBIC, expand=True)

            except Exception as osd_error:
                st.warning(f"  > Strona {i+1}: Błąd podczas wykrywania orientacji: {osd_error}. Próbuję OCR na oryginalnym obrazie.")
                # W razie błędu OSD, kontynuuj z oryginalnym obrazem
                rotated_image = img

            # 3. Wykonaj OCR na (potencjalnie obróconym) obrazie
            #st.write(f"  > Strona {i+1}: Wykonywanie OCR...")
            # Konwersja obrazu PIL do formatu akceptowanego przez Tesseract
            # Można rozważyć usunięcie "+eng" jeśli OSD poprawnie wykrywa skrypt,
            # ale pozostawienie go może pomóc w trudniejszych przypadkach.
            # Używamy domyślnego PSM (3), które jest zazwyczaj dobre po korekcji orientacji.
            text = pytesseract.image_to_string(rotated_image, lang=lang_code+"+eng")
            full_text += text + "\n\n"  # Dodaj separator między stronami

        return full_text.strip()
    except Exception as e:
        st.error(f"Błąd podczas OCR: {e}")
        return None

def translate_text_stream(text_to_translate, source_lang_name, target_lang_llm):
    """Wysyła tekst do OpenRouter API i streamuje tłumaczenie."""
    if not OPENROUTER_API_KEY or OPENROUTER_API_KEY == "sk-or-v1-...":
        st.error("Klucz API OpenRouter nie został ustawiony. Edytuj plik app.py.")
        return None

    client = OpenAI(
        #base_url="https://openrouter.ai/api/v1",
        base_url="http://172.17.0.1:8000/v1",
        api_key=OPENROUTER_API_KEY,
    )

    prompt = f"Translate below text to {target_lang_llm}:\n\n{text_to_translate}\n\nTranslation to {target_lang_llm} in MARKDOWN FORMAT:"

    # Pobierz odpowiedni komunikat systemowy
    system_message_content = SYSTEM_MESSAGES.get(target_lang_llm, SYSTEM_MESSAGES.get("default"))
    if not system_message_content:
        st.warning("Nie znaleziono domyślnego komunikatu systemowego. Używam pustego.")
        system_message_content = ""

    try:
        #print(system_message_content)
        #print(prompt)
        stream = client.chat.completions.create(
            #model="google/gemma-3-27b-it",
            model="/model",
            messages=[
                {"role": "system", "content": system_message_content},
                {"role": "user", "content": prompt},
            ],
            stream=True,
            temperature=0.1,
            top_p=0.95,
        )
        return stream
    except Exception as e:
        st.error(f"Błąd podczas komunikacji z OpenRouter API: {e}")
        return None

# --- Funkcja opakowująca strumień --- ZMODYFIKOWANA
def wrap_stream_for_markdown(stream):
    """Generator opakowujący strumień OpenAI, usuwający potencjalny
       blok kodu markdown (```markdown\n) na początku odpowiedzi.
    """
    first_chunk_processed = False
    buffer = ""
    leading_sequence = "```markdown\n"
    sequence_removed = False

    for chunk in stream:
        # Sprawdź, czy chunk ma oczekiwaną strukturę i zawartość
        try:
            content = chunk.choices[0].delta.content
            #logging.info(f"LLM Raw Stream Content: {repr(content)}") # Zmieniono log prefix
        except (AttributeError, IndexError, TypeError):
            #logging.warning(f"Problematic chunk structure: {chunk}")
            yield chunk # Przekaż problematyczny chunk dalej
            continue

        if content is None:
            content = "" # Traktuj None jako pusty string

        if not first_chunk_processed and not sequence_removed:
            # Buforuj, dopóki nie zbierzemy wystarczająco dużo, by sprawdzić sekwencję
            buffer += content
            #logging.debug(f"Buffering: {repr(buffer)}")

            # Sprawdź, czy bufor zaczyna się od sekwencji (ignorując białe znaki na początku)
            stripped_buffer = buffer.lstrip()
            if stripped_buffer.startswith(leading_sequence):
                # Znaleziono sekwencję, usuń ją
                buffer = stripped_buffer[len(leading_sequence):]
                #logging.info(f"Removed leading sequence. Remaining buffer: {repr(buffer)}")
                sequence_removed = True
                first_chunk_processed = True # Pierwszy "znaczący" fragment przetworzony

                # Jeśli coś zostało w buforze po usunięciu, zwróć to jako pierwszy chunk
                if buffer:
                    # Stwórz nowy chunk z pozostałością bufora
                    # To jest uproszczenie, zakładamy że struktura chunk jest podobna
                    # Może wymagać dostosowania jeśli API zwróci inną strukturę
                    try:
                         chunk.choices[0].delta.content = buffer
                         #logging.info(f"Yielding modified first chunk: {repr(buffer)}")
                         yield chunk
                    except Exception as e:
                         #logging.error(f"Error modifying chunk: {e}")
                         # W razie błędu zwróć oryginalny chunk (może być pusty)
                         yield chunk
                buffer = "" # Wyczyść bufor

            elif len(buffer) > len(leading_sequence) + 5: # Daj trochę zapasu
                # Jeśli zebraliśmy wystarczająco dużo i sekwencji nie ma, przestajemy buforować
                #logging.info(f"Leading sequence not found. Yielding buffered content: {repr(buffer)}")
                first_chunk_processed = True
                # Zwróć cały bufor jako pierwszy chunk
                try:
                     chunk.choices[0].delta.content = buffer
                     yield chunk
                except Exception as e:
                     #logging.error(f"Error modifying chunk: {e}")
                     yield chunk # Zwróć oryginalny
                buffer = ""
            # Jeśli bufor jest krótszy niż sekwencja, kontynuuj buforowanie

        else: # Pierwszy chunk przetworzony lub sekwencja już usunięta
            # Po prostu zwróć oryginalny chunk (jeśli ma zawartość)
            if content:
                 #logging.debug(f"Yielding subsequent chunk: {repr(content)}")
                 yield chunk
            elif chunk.choices and chunk.choices[0].finish_reason: # Zwróć chunk kończący
                 #logging.debug("Yielding final chunk.")
                 yield chunk
            # Ignoruj puste chunki pośrednie

# --- Funkcje eksportu ---

def markdown_to_docx(markdown_text, output_filename="translation_export.docx"):
    """Konwertuje tekst w formacie Markdown na plik DOCX."""
    # Konwertuj Markdown na HTML
    html = markdown.markdown(markdown_text)
    
    # Tworzenie dokumentu
    doc = Document()
    
    # Dodanie stylu do dokumentu
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Sparsuj HTML, aby ekstrakcja tekstu była łatwiejsza
    soup = BeautifulSoup(html, 'html.parser')
    
    # Przetwarzamy każdy element HTML
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li', 'blockquote']):
        if element.name.startswith('h'):
            # Nagłówki
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)
        elif element.name == 'p':
            # Paragraf tekstu
            p = doc.add_paragraph(element.get_text())
            # Obsługa podstawowego formatowania
            for child in element.children:
                if child.name == 'strong' or child.name == 'b':
                    for run in p.runs:
                        run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    for run in p.runs:
                        run.italic = True
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == 'blockquote':
            doc.add_paragraph(element.get_text()).style = 'Quote'
    
    # Zapisz dokument do pamięci
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

def markdown_to_pdf(markdown_text, output_filename="translation_export.pdf"):
    """Konwertuje tekst w formacie Markdown na plik PDF."""
    # Konwertuj Markdown na HTML
    html_content = markdown.markdown(markdown_text)
    
    # Dodanie podstawowych stylów dla lepszego wyglądu
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.5; margin: 2cm; }}
            h1, h2, h3, h4, h5, h6 {{ color: #333; margin-top: 1em; }}
            p {{ margin: 0.5em 0; }}
            blockquote {{ border-left: 3px solid #ccc; padding-left: 1em; color: #666; }}
            ul, ol {{ margin: 0.5em 0; padding-left: 2em; }}
            code {{ background: #f4f4f4; padding: 0.2em 0.4em; border-radius: 3px; }}
            pre {{ background: #f4f4f4; padding: 1em; border-radius: 5px; overflow-x: auto; }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    # Konwersja HTML na PDF za pomocą WeasyPrint
    pdf_bytes = io.BytesIO()
    HTML(string=styled_html).write_pdf(pdf_bytes)
    pdf_bytes.seek(0)
    
    return pdf_bytes

def get_download_link(file_bytes, filename, text):
    """Generuje link do pobrania pliku."""
    b64 = base64.b64encode(file_bytes.read()).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{text}</a>'
    return href

# --- Interfejs Użytkownika Streamlit ---

st.set_page_config(layout="wide") # Użyj szerokiego layoutu

# Dodaj niestandardowy CSS, aby zmniejszyć górny padding
st.markdown("""
<style>
    /* Celuje w główny kontener bloku */
    .block-container {
        padding-top: 1rem !important; /* Możesz dostosować tę wartość (np. 0rem) */
    }
    /* Czasami potrzeba bardziej specyficznego selektora */
    /* div[data-testid="stAppViewBlockContainer"] {
        padding-top: 1rem !important;
    }*/
</style>
""", unsafe_allow_html=True)

st.title("📄 PDF Translator Demo (OCR + LLM)")

# --- Pasek Boczny (Sidebar) ---
with st.sidebar:
    st.header("Ustawienia Tłumaczenia")

    uploaded_file = st.file_uploader(
        "1. Załaduj plik PDF (bez warstwy tekstowej)", type="pdf"
    )

    ocr_lang_name = st.selectbox(
        "3. Wybierz język źródłowy (OCR):",
        options=list(LANGUAGES.keys()),
        index=0,  # Domyślnie Angielski
        key="ocr_lang"
    )
    ocr_lang_code = LANGUAGES[ocr_lang_name][0]

    target_lang_name = st.selectbox(
        "4. Wybierz język docelowy (Tłumaczenie):",
        options=list(LANGUAGES.keys()),
        index=2,  # Domyślnie Polski
        key="target_lang"
    )
    target_lang_llm = LANGUAGES[target_lang_name][1]
    # Pole do wyboru stron - dodane tutaj
    page_selection_str = st.text_input(
        "2. Wybierz strony (np. 1, 3, 5-7, puste = wszystkie):",
        key="page_selection" # Klucz dla stanu sesji
    )
    # Użyj kolumn dla przycisków
    col1_sidebar, col2_sidebar = st.columns(2)
    with col1_sidebar:
        translate_button = st.button(
            "🚀 Przetłumacz", disabled=not uploaded_file, use_container_width=True, key="translate_btn"
        )
    with col2_sidebar:
        # Dodajemy przycisk Reset
        reset_button = st.button("🔄 Resetuj Stan", use_container_width=True)

    st.markdown("---")

    # Miejsce na komunikaty zwrotne
    feedback_placeholder = st.empty()

# --- Reset Logic ---
# Umieść ten blok POZA `with st.sidebar:` ale PRZED główną logiką przetwarzania
if reset_button:
    # Wyczyszczenie zmiennych stanu sesji
    keys_to_reset = [
        'images', 'ocr_text', 'translation_stream', 'error_message',
        'success_message', 'total_pages_in_doc', 'selected_page_numbers'
        # Usunięto 'page_selection' - chcemy zachować wybór stron
    ]
    for key in keys_to_reset:
        if key in st.session_state:
            st.session_state[key] = None
    # Wyczyszczenie miejsca na komunikaty
    feedback_placeholder.empty()
    # Odświeżenie aplikacji
    st.rerun() # Użyj st.rerun() zamiast st.experimental_rerun() w nowszych wersjach Streamlit

# --- Główny Obszar (Podział na Kolumny) ---

# Inicjalizacja zmiennych stanu poza głównym blokiem if
if 'images' not in st.session_state:
    st.session_state.images = None
if 'ocr_text' not in st.session_state:
    st.session_state.ocr_text = None
if 'translation_stream' not in st.session_state:
    st.session_state.translation_stream = None
if 'error_message' not in st.session_state:
    st.session_state.error_message = None
if 'success_message' not in st.session_state:
    st.session_state.success_message = None
# Dodane zmienne stanu
if 'total_pages_in_doc' not in st.session_state:
    st.session_state.total_pages_in_doc = None
if 'selected_page_numbers' not in st.session_state:
     st.session_state.selected_page_numbers = None
# Zmienne stanu dla eksportu
if 'full_translation' not in st.session_state:
    st.session_state.full_translation = None
if 'export_docx_link' not in st.session_state:
    st.session_state.export_docx_link = None
if 'export_pdf_link' not in st.session_state:
    st.session_state.export_pdf_link = None
if 'translation_displayed' not in st.session_state:
    st.session_state.translation_displayed = False

# Funkcje dla obsługi przycisków eksportu
def generate_docx():
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"translation_export_{timestamp}.docx"
        docx_bytes = markdown_to_docx(st.session_state.full_translation, docx_filename)
        
        # Zapisz link w stanie sesji
        b64 = base64.b64encode(docx_bytes.getvalue()).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="{docx_filename}">📥 Pobierz plik DOCX</a>'
        st.session_state.export_docx_link = href
    except Exception as e:
        st.error(f"Błąd podczas generowania DOCX: {e}")
        st.session_state.export_docx_link = None

def generate_pdf():
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"translation_export_{timestamp}.pdf"
        pdf_bytes = markdown_to_pdf(st.session_state.full_translation, pdf_filename)
        
        # Zapisz link w stanie sesji
        b64 = base64.b64encode(pdf_bytes.getvalue()).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="{pdf_filename}">📥 Pobierz plik PDF</a>'
        st.session_state.export_pdf_link = href
    except Exception as e:
        st.error(f"Błąd podczas generowania PDF: {e}")
        st.session_state.export_pdf_link = None

col1, col2 = st.columns(2)

if translate_button and uploaded_file is not None:
    pdf_bytes = uploaded_file.getvalue()
    # Resetuj stan przy nowym przetwarzaniu, ale zachowaj input użytkownika
    st.session_state.images = None
    st.session_state.ocr_text = None
    st.session_state.translation_stream = None
    st.session_state.error_message = None
    st.session_state.success_message = None
    st.session_state.total_pages_in_doc = None
    st.session_state.selected_page_numbers = None

    # Krok 0: Odczytaj liczbę stron i sparsuj wybór użytkownika
    try:
        # Używamy funkcji extract_images... do otwarcia i pobrania liczby stron
        # Przekazujemy None jako selected_pages, żeby tylko odczytać liczbę stron
        _, total_pages = extract_images_from_pdf(pdf_bytes, selected_pages=None)
        st.session_state.total_pages_in_doc = total_pages
        if total_pages == 0 and not st.session_state.error_message: # Sprawdź czy błąd nie wystąpił już w extract_images...
             st.session_state.error_message = "Nie udało się odczytać liczby stron z pliku PDF."

        if not st.session_state.error_message:
            # Parsuj wybór stron użytkownika (użyj wartości z st.session_state, aby zachować ją między uruchomieniami)
            current_page_selection = st.session_state.get("page_selection", "") # Pobierz z stanu sesji
            selected_pages = parse_page_numbers(current_page_selection, total_pages)

            if selected_pages is None:
                # Błąd parsowania został już wyświetlony przez parse_page_numbers
                st.session_state.error_message = "Popraw wybór stron w pasku bocznym." # Dodatkowy komunikat
            else:
                st.session_state.selected_page_numbers = selected_pages
                if not selected_pages: # Jeśli zwrócono pustą listę (np. po walidacji w parse..)
                     st.session_state.error_message = "Nie wybrano żadnych prawidłowych stron do przetworzenia."


    except Exception as e:
        st.session_state.error_message = f"Błąd podczas wstępnego przetwarzania PDF: {e}"

    # Krok 1: Ekstrakcja obrazów (tylko jeśli nie było błędów wcześniej)
    if not st.session_state.error_message:
        with feedback_placeholder.status(f"Ekstrahowanie {len(st.session_state.selected_page_numbers)}/{st.session_state.total_pages_in_doc} stron z PDF...", expanded=True) as status:
            # Użyj sparsowanych i zwalidowanych numerów stron
            images, _ = extract_images_from_pdf(pdf_bytes, selected_pages=st.session_state.selected_page_numbers)
            st.session_state.images = images # Zapisz obrazy w stanie sesji

            if st.session_state.images is None: # Sprawdź, czy ekstrakcja się powiodła (extract_images_from_pdf zwraca None w razie błędu)
                # Błąd powinien być już ustawiony wewnątrz extract_images_from_pdf
                if not st.session_state.error_message: # Na wszelki wypadek
                     st.session_state.error_message = "Nieznany błąd podczas ekstrakcji obrazów."
                status.update(label="Błąd ekstrakcji!", state="error", expanded=True)
            elif not st.session_state.images: # Pusta lista obrazów (np. jeśli wybrano 0 stron)
                 # Komunikat ostrzegawczy powinien pojawić się w extract lub parse
                 if not st.session_state.error_message and not st.session_state.success_message: # Jeśli nie ma już komunikatu
                     st.warning("Nie wybrano żadnych stron do przetworzenia.")
                 status.update(label="Brak stron do ekstrakcji.", state="complete", expanded=False)
            else:
                status.update(label=f"Wyekstrahowano {len(st.session_state.images)} obrazów.", state="complete", expanded=False)

    # Krok 2: OCR
    # Sprawdzamy czy są obrazy i nie ma błędu
    if st.session_state.images and not st.session_state.error_message:
         with feedback_placeholder.status(f"Wykonywanie OCR ({ocr_lang_name}) dla {len(st.session_state.images)} stron...", expanded=True) as status:
            st.session_state.ocr_text = perform_ocr(st.session_state.images, ocr_lang_code)
            # Sprawdzamy, czy OCR zwrócił tekst (perform_ocr zwraca "" dla pustej listy obrazów)
            if st.session_state.ocr_text is None: # Błąd w perform_ocr
                # Błąd powinien być ustawiony w perform_ocr
                 if not st.session_state.error_message:
                     st.session_state.error_message = "Nie udało się wykonać OCR na pliku."
                 status.update(label="Błąd OCR!", state="error", expanded=True)
            elif not st.session_state.ocr_text and st.session_state.images: # Jeśli były obrazy, ale OCR nic nie zwrócił
                 st.warning("OCR nie rozpoznał tekstu na wybranych stronach.")
                 status.update(label="OCR nie znalazł tekstu.", state="complete", expanded=False)
            else:
                 status.update(label="OCR zakończony.", state="complete", expanded=False)

    # Krok 3: Tłumaczenie
    # Sprawdzamy, czy jest tekst OCR i nie ma błędu
    if st.session_state.ocr_text and not st.session_state.error_message:
        with feedback_placeholder.status(f"Tłumaczenie z {ocr_lang_name} na {target_lang_name}...", expanded=True) as status:
            try:
                st.session_state.translation_stream = translate_text_stream(
                    st.session_state.ocr_text, ocr_lang_name, target_lang_llm
                )
                if not st.session_state.translation_stream:
                     st.session_state.error_message = "Nie udało się rozpocząć procesu tłumaczenia (problem z API?)."
                     status.update(label="Błąd inicjalizacji tłumaczenia!", state="error", expanded=True)
                else:
                    # Symulacja zakończenia, bo stream będzie w kolumnie
                    st.session_state.success_message = "Przetwarzanie rozpoczęte." # Zmieniono komunikat
                    status.update(label="Tłumaczenie rozpoczęte.", state="complete", expanded=False) # Zmieniono status

            except Exception as e:
                 st.session_state.error_message = f"Błąd podczas komunikacji z API: {e}"
                 status.update(label="Błąd API!", state="error", expanded=True)

    # Wyświetlanie końcowych komunikatów w sidebarze
    if st.session_state.error_message:
        feedback_placeholder.error(st.session_state.error_message)
    elif st.session_state.success_message and not st.session_state.translation_stream: # Dodano warunek, by nie nadpisywać statusu stream
         feedback_placeholder.success(st.session_state.success_message)
         # Sukces jest teraz implikowany przez obecność strumienia, komunikat wyświetlany w trakcie
         #pass # Usunięto pass, aby komunikat sukcesu mógł się pojawić jeśli nie ma streamu


# --- Wyświetlanie Wyników w Kolumnach ---

# --- Kolumna Lewa: Oryginalny PDF (jako obrazy) ---
with col1:
    st.subheader("📄 Oryginalny Dokument (Strony)")
    if st.session_state.images:
        pdf_container = st.container(height=700)
        with pdf_container:
            for i, img in enumerate(st.session_state.images):
                # Zmiana use_column_width na use_container_width
                st.image(img, caption=f"Strona {i+1}", use_container_width=True)
    elif uploaded_file and not st.session_state.images and not st.session_state.error_message:
         # Ten przypadek jest już obsłużony przez error message w sidebarze
         pass
    elif not uploaded_file:
        st.info("Załaduj plik PDF w pasku bocznym.")

# --- Kolumna Prawa: OCR i Tłumaczenie ---
with col2:
    st.subheader("📝 Wyniki Przetwarzania")

    if st.session_state.ocr_text:
        with st.expander("🔍 Pokaż tekst rozpoznany przez OCR", expanded=False):
            st.text_area("Tekst z OCR", st.session_state.ocr_text, height=200, disabled=True, key="ocr_output")
    elif uploaded_file and not st.session_state.images and not st.session_state.error_message:
        pass # Obsłużone w sidebarze
    elif uploaded_file and st.session_state.images and not st.session_state.ocr_text and not st.session_state.error_message:
         pass # Obsłużone w sidebarze


    if st.session_state.translation_stream and not st.session_state.translation_displayed:
        st.subheader("✅ Wynik Tłumaczenia:")
        try:
            output_container = st.container(height=570) # Kontener dla tłumaczenia
            with output_container:
                # Streamlit wymaga, aby `write_stream` był poza `st.status`
                # Ponieważ `translation_stream` jest generatorem, musi być konsumowany tutaj.
                # Użyj funkcji opakowującej!
                wrapped_stream = wrap_stream_for_markdown(st.session_state.translation_stream)
                full_response = st.write_stream(wrapped_stream)
                # Zapisz pełną odpowiedź, aby móc ją eksportować
                st.session_state.full_translation = full_response
                # Ustaw flagę, że tłumaczenie zostało wyświetlone
                st.session_state.translation_displayed = True

            # Wyświetl sukces PO zakończeniu streamowania
            feedback_placeholder.success("Tłumaczenie zakończone!")
            st.session_state.success_message = "Tłumaczenie zakończone!" # Ustawienie flagi sukcesu po zakończeniu
        
        except Exception as e:
            # Błąd podczas samego streamowania
            error_msg = f"Błąd podczas streamowania odpowiedzi: {e}"
            st.error(error_msg) # Błąd wyświetlany bezpośrednio w kolumnie
            feedback_placeholder.error(error_msg) # Oraz w sidebarze
            st.session_state.error_message = error_msg # Zapisz błąd
    
    # Jeśli tłumaczenie już zostało wyświetlone wcześniej (ale nie mamy linków do plików)
    elif st.session_state.full_translation is not None and st.session_state.translation_displayed:
        st.subheader("✅ Wynik Tłumaczenia:")
        translation_container = st.container(height=570)
        with translation_container:
            # Wyświetl zapisane tłumaczenie ponownie, aby było widoczne po kliknięciu przycisków
            st.markdown(st.session_state.full_translation)
    
    # Sprawdź czy jest zapisane tłumaczenie i czy już się zakończyło
    if st.session_state.full_translation is not None and st.session_state.translation_displayed:
        # Dodaj przyciski eksportu PO zakończeniu tłumaczenia
        st.subheader("📥 Eksport Tłumaczenia:")
        export_cols = st.columns(2)
        
        with export_cols[0]:
            # Przycisk do pobrania jako DOCX
            if st.button("📄 Eksportuj do Word (DOCX)", key="export_docx"):
                generate_docx() # Wywołaj funkcję generowania DOCX
                st.rerun() # Odśwież, aby pokazać link
            
            # Wyświetl link do pobrania DOCX, jeśli został wygenerowany
            if st.session_state.export_docx_link:
                st.markdown(st.session_state.export_docx_link, unsafe_allow_html=True)
        
        with export_cols[1]:
            # Przycisk do pobrania jako PDF
            if st.button("📄 Eksportuj do PDF", key="export_pdf"):
                generate_pdf() # Wywołaj funkcję generowania PDF
                st.rerun() # Odśwież, aby pokazać link
            
            # Wyświetl link do pobrania PDF, jeśli został wygenerowany
            if st.session_state.export_pdf_link:
                st.markdown(st.session_state.export_pdf_link, unsafe_allow_html=True)
    
    elif uploaded_file and st.session_state.ocr_text and not st.session_state.translation_stream and not st.session_state.error_message:
        pass # Obsłużone w sidebarze
    elif not uploaded_file:
         st.info("Wyniki pojawią się tutaj po przetworzeniu.")

